from docx import Document
from docx.shared import Inches, Cm, Pt



# ask report identification and setup - verify input and create/retreive doc
from check_language_selection import select_destination_language
from report_settings import get_report_language, set_report_language, set_report_candidate_name, set_report_document_code
import get_doc_id
import get_report_CF

# Ask user for the document code of the report (e.g. 'AB-123')

CF_report_number = input('Provide document code: ')

cf_report_docx = Document()

run = cf_report_docx.add_heading('Angst en Moed rapportage', 0).add_run()
font = run.font
p = cf_report_docx.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

font.name = 'Calibri'
font.size = Pt(12) Pt
font.name = 'Calibri'
font.size = Pt(12)

cf_report_docx.add_heading('Heading, level 1', level=1)
cf_report_docx.add_paragraph('Intense quote', style='Intense Quote')

cf_report_docx.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
cf_report_docx.add_paragraph(
    'first item in ordered list', style='List Number'
)

cf_report_docx.add_picture('monty-truth.png', width=Cm(3.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = cf_report_docx.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

cf_report_docx.add_page_break()

report_file_name = 'CF_report_demo'+str(CF_report_number)+'.docx'
cf_report_docx.save(report_file_name)